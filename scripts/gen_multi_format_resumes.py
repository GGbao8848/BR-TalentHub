"""生成多格式测试简历（md/docx/doc/pdf/png）并通过 API 上传。"""
import io
import urllib.request
import uuid
from pathlib import Path

from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas
from docx import Document

BASE_DIR = Path(__file__).resolve().parent.parent
API = "http://127.0.0.1:18540/api/resumes/upload"

NAME = "多格式测试"
PHONE = "13900001111"
POSITION = "研发工程师"
SCHOOL = "北京大学"


def upload(filename: str, data: bytes):
    boundary = uuid.uuid4().hex
    body = io.BytesIO()
    fields = {
        "name": NAME,
        "phone": PHONE,
        "position": POSITION,
        "school": SCHOOL,
    }
    for k, v in fields.items():
        body.write(f"--{boundary}\r\nContent-Disposition: form-data; name=\"{k}\"\r\n\r\n{v}\r\n".encode())
    body.write(
        f"--{boundary}\r\nContent-Disposition: form-data; name=\"file\"; filename=\"{filename}\"\r\n"
        f"Content-Type: application/octet-stream\r\n\r\n".encode()
    )
    body.write(data)
    body.write(f"\r\n--{boundary}--\r\n".encode())
    req = urllib.request.Request(
        API, data=body.getvalue(),
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
    )
    try:
        resp = urllib.request.urlopen(req, timeout=30)
        print(f"  ✓ {filename}: {resp.status} {resp.read().decode()[:80]}")
    except urllib.error.HTTPError as e:
        print(f"  ✗ {filename}: {e.code} {e.read().decode()[:120]}")


def make_md() -> bytes:
    content = (
        "# 张三 · 研发工程师\n\n"
        "## 基本信息\n"
        "- 姓名：张三\n- 电话：13900001111\n- 院校：北京大学\n\n"
        "## 技能\n"
        "- Python / FastAPI\n- Vue3 / Naive UI\n- MySQL / Redis\n\n"
        "## 项目经验\n"
        "参与了多个中大型系统开发，负责后端接口设计与性能优化。\n\n"
        "> 这是一份 **Markdown** 格式的简历，用于测试查看器对 .md 文件的渲染。\n"
    )
    return content.encode()


def make_docx() -> bytes:
    doc = Document()
    doc.add_heading("李四 · 研发工程师", 0)
    doc.add_heading("基本信息", level=1)
    doc.add_paragraph("姓名：李四")
    doc.add_paragraph("电话：13900001111")
    doc.add_paragraph("院校：北京大学")
    doc.add_heading("工作经历", level=1)
    doc.add_paragraph("2020-2024 在某互联网公司担任后端开发，负责核心业务模块。")
    doc.add_paragraph("熟悉高并发场景下的系统设计与调优。")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_doc() -> bytes:
    # .doc 为旧版二进制格式，用富文本 RTF 内容（Word 可打开）作为近似
    rtf = (
        "{\\rtf1\\ansi\\deff0 {\\fonttbl {\\f0 Microsoft YaHei;}}\n"
        "\\f0\\fs24 王五 · 研发工程师\\par\n"
        "\\b 基本信息\\b0\\par\n"
        "姓名：王五\\par 电话：13900001111\\par 院校：北京大学\\par\n"
        "\\b 工作经历\\b0\\par\n"
        "五年后端开发经验，负责过多个大型项目。\\par\n"
        "}"
    )
    return rtf.encode("utf-8")


def make_pdf() -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    c.setFont("STSong-Light", 22)
    c.drawString(80, 780, "赵六 · 研发工程师")
    c.setFont("STSong-Light", 13)
    y = 740
    for line in ["姓名：赵六", "电话：13900001111", "院校：北京大学", "", "工作经历：", "参与多个系统开发，负责后端架构设计。"]:
        c.drawString(90, y, line)
        y -= 22
    c.showPage()
    c.save()
    return buf.getvalue()


def make_png() -> bytes:
    img = Image.new("RGB", (800, 1000), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, 799, 120], fill="#2563eb")
    draw.text((40, 40), "孙七 · 研发工程师", fill="white")
    draw.text((40, 160), "姓名：孙七", fill="black")
    draw.text((40, 200), "电话：13900001111", fill="black")
    draw.text((40, 240), "院校：北京大学", fill="black")
    draw.text((40, 300), "工作经历：", fill="black")
    draw.text((40, 340), "负责前端界面开发，熟悉 Vue 与组件库。", fill="black")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


if __name__ == "__main__":
    jobs = [
        ("测试简历_研发工程师.md", make_md()),
        ("测试简历_研发工程师.docx", make_docx()),
        ("测试简历_研发工程师.doc", make_doc()),
        ("测试简历_研发工程师.pdf", make_pdf()),
        ("测试简历_研发工程师.png", make_png()),
    ]
    print(f"上传到 {API}")
    for fn, data in jobs:
        upload(fn, data)
    print("完成")
